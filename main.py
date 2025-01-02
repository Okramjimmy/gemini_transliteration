from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
import google.generativeai as genai
import os
from docx import Document
import json
from dotenv import load_dotenv
import io
import re
load_dotenv()  # Load variables from .env file

app = FastAPI()

# Configure your Gemini API key
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
INPUT_LANG = os.getenv("INPUT_LANG")
OUTPUT_LANG = os.getenv("OUTPUT_LANG")
LLM_MODEL = os.getenv("LLM_MODEL")


if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY not found in .env file.")
if not INPUT_LANG:
    raise ValueError("INPUT_LANG not found in .env file.")
if not OUTPUT_LANG:
    raise ValueError("OUTPUT_LANG not found in .env file.")
if not LLM_MODEL:
    raise ValueError("LLM_MODEL not found in .env file.")

genai.configure(api_key=GEMINI_API_KEY)

# Use gemini 1.5 pro
model = genai.GenerativeModel(model_name=LLM_MODEL)


# Pydantic model for successful response
class TransliterationResponse(BaseModel):
    original_text: str
    transliterated_text: str

# Pydantic model for successful NER response
class NERResponse(BaseModel):
    entities: dict

# Pydantic model for Error response
class ErrorResponse(BaseModel):
    error:str

class WelcomeResponse(BaseModel):
    message: str

@app.post("/", response_model=WelcomeResponse)
async def welcome():
    """
    Root endpoint for the API.
    Returns a welcome message with a link to the API documentation.
    """
    return WelcomeResponse(message="Welcome to the AI Transliteration and NER API! Please go to `/docs` for the API documentation.")

@app.post("/transliterate/text", response_model=TransliterationResponse | ErrorResponse, responses={400: {"model": ErrorResponse}})
async def transliterate_text(text: str = Form(...)):
    """
    Transliterates text from English to Assamese using Gemini 1.5 Flash.
    Handles plain text input.
    Returns transliterated text in JSON format.
    """
    try:
        if not text:
            raise HTTPException(status_code=400, detail="Please provide text input.")

        prompt = f"Transliterate the following text from {INPUT_LANG} to {OUTPUT_LANG} script: '{text}'. Please respond with only transliterated {OUTPUT_LANG} script. Do not provide any extra text and dont translate the text"
        response = model.generate_content(prompt)
        transliterated_text = response.text.strip()

        return TransliterationResponse(original_text=text, transliterated_text=transliterated_text)
    except Exception as e:
        return ErrorResponse(error=str(e))

@app.post("/transliterate/file", response_model=TransliterationResponse | ErrorResponse, responses={400: {"model": ErrorResponse}})
async def transliterate_file(file: UploadFile = File(...)):
    """
    Transliterates text from a docx file from English to Assamese using Gemini 1.5 Flash.
    Handles .docx file uploads.
    Returns transliterated text in JSON format.
    """
    try:
        if not file:
            raise HTTPException(status_code=400, detail="Please provide a file.")

        if file.content_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            raise HTTPException(status_code=400, detail="Invalid file type. Please upload a .docx file.")

        contents = await file.read()
        doc = Document(io.BytesIO(contents))
        input_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        prompt = f"Transliterate the following text from {INPUT_LANG} to {OUTPUT_LANG} script: '{text}'. Please respond with only transliterated {OUTPUT_LANG} script. Do not provide any extra text and dont translate the text"
        response = model.generate_content(prompt)
        transliterated_text = response.text.strip()

        return TransliterationResponse(original_text=input_text, transliterated_text=transliterated_text)
    except Exception as e:
        return ErrorResponse(error=str(e))
@app.post("/ner", response_model=NERResponse | ErrorResponse, responses={400: {"model": ErrorResponse}})
async def extract_and_transliterate_ner(file: UploadFile = File(...)):
    """
    Extracts Named Entities from a docx file and transliterates them to Assamese.
    Returns a key-value JSON where the keys are the original English entities
    and the values are their Assamese transliterations.
    """
    try:
        if not file:
            raise HTTPException(status_code=400, detail="Please provide a file.")

        if file.content_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            raise HTTPException(status_code=400, detail="Invalid file type. Please upload a .docx file.")

        contents = await file.read()
        doc = Document(io.BytesIO(contents))
        input_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # ner_prompt = f"""Extract named entities (people, organizations, locations, dates, etc) from the following text and transliterate each entity to Assamese.
        # Return a json output where each named entity is a key and its corresponding transliterated assamese value.
        # Text: '{input_text}'. Please return json only"""
        ner_prompt = f"""
            Extract named entities (people, organizations, locations, dates, legal terms, legal abbreviations, act names, institute names, law journal names, medical terms, proper nouns, foreign language terms like Latin words, book names, and vehicle numbers) from the following text. 
            For each extracted entity, transliterate it into {OUTPUT_LANG} script. 
            Return a JSON object where each named entity is a key, and its corresponding transliterated {OUTPUT_LANG} value is the value.

            Text: '{input_text}'
            Please return only the JSON object.
            """


        ner_response = model.generate_content(ner_prompt)
        ner_text = ner_response.text.strip()

        # Attempt to parse the response, try to catch cases where it is not a json
        try:
            entities = json.loads(ner_text)
        except json.JSONDecodeError:
            #if json is not extracted from the output, try to extract it using regex
            pattern = r"```(?:json)?(.*?)```"
            match = re.search(pattern, ner_text, re.DOTALL)
            if match:
                try:
                    entities = json.loads(match.group(1).strip())
                except json.JSONDecodeError:
                    raise HTTPException(status_code=400, detail=f"Could not extract json for NER entities : {ner_text}")
            else:
                raise HTTPException(status_code=400, detail=f"Could not extract json for NER entities : {ner_text}")
        
        return NERResponse(entities=entities)

    except Exception as e:
        return ErrorResponse(error=str(e))

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)